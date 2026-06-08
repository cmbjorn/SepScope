"""
Comprehensive Word (.docx) scoping report generator for SepScope.

Structure
---------
  Cover      — project / vessel header, revision block, issued-for stamp
  A          — Design conditions (operating & design P/T, hydro test)
  B          — Process fluids (gas & liquid properties, nozzle stream table)
  C          — Separator sizing (API 12J screening table)
  C.1        — Turndown Analysis
  C.2        — LDV — Liquid Design Volume (when applicable)
  C.3        — Internals — Mechanical Loads (LDV Startup Surge) (when applicable)
  D          — Mechanical design (geometry, calculated thicknesses, material)
  D.1        — Overall height & mounting (top of vessel → bottom of saddle feet)
  D.2        — Weight estimate (dry / operating / hydro)
  D.3        — Shell thickness calculation narrative
  D.4        — Head thickness calculation narrative
  D.5        — Internal Lining / Surface Treatment (when applicable)
  E          — Liquid levels & volumes
  F          — Internals
  F.1        — Inlet Device Sizing (API 12J §5.3)
  G          — Nozzle schedule
  G.1        — Endcap nozzle detail (zone, edge clearances, head comparison)
  H          — Engineering findings & notes

All dimensions in mm, pressures in barg/MPa, volumes in m³ and L.
"""
from __future__ import annotations
import io
import math
from datetime import date as _date
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from engines.nozzle_reinforcement import suggest_schedule_upgrade
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────────────
_C_DARK   = "1E3A5F"   # main header blue
_C_MID    = "2D5F8A"   # sub-panel header
_C_LIGHT  = "DCE4EF"   # header row fill
_C_ALT    = "F4F6FA"   # alternating row fill
_C_OK     = "D1FAE5"   # green cell
_C_WARN   = "FEF3C7"   # amber cell
_C_FAIL   = "FEE2E2"   # red cell
_C_WHITE  = "FFFFFF"
_C_BORDER = "B0B8C4"

_RGB_DARK   = RGBColor(0x1E, 0x3A, 0x5F)
_RGB_MID    = RGBColor(0x2D, 0x5F, 0x8A)
_RGB_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_RGB_OK     = RGBColor(0x16, 0x65, 0x34)
_RGB_WARN   = RGBColor(0x92, 0x40, 0x0E)
_RGB_FAIL   = RGBColor(0xDC, 0x26, 0x26)
_RGB_GREY   = RGBColor(0x64, 0x74, 0x8B)

# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, hex6: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd to avoid duplicates
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)


def _set_table_borders(table, color: str = _C_BORDER, sz: str = "4") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    # replace old tblBorders if present
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(w * 567)))   # 1 cm ≈ 567 twips
            tcW.set(qn("w:type"), "dxa")
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcPr.append(tcW)


def _no_space_para(para) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


# ── Document setup ────────────────────────────────────────────────────────────

def _setup_document(doc: Document) -> None:
    """A4, 2.5 cm margins, default font."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.2))

    # Normal style baseline
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    # Add footer with page number
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = _RGB_GREY
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    run._r.append(instrText)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)
    run2 = fp.add_run(" / ")
    run2.font.size = Pt(8)
    run2.font.color.rgb = _RGB_GREY
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    run2._r.append(fld3)
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = "NUMPAGES"
    run2._r.append(instrText2)
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    run2._r.append(fld4)


# ── Paragraph / heading helpers ───────────────────────────────────────────────

def _section_heading(doc: Document, letter: str, title: str) -> None:
    """Dark-blue banner heading: 'A — Design Conditions'."""
    p = doc.add_paragraph()
    _no_space_para(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f"  {letter} — {title}  ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _RGB_WHITE
    # blue shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _C_DARK)
    pPr.append(shd)


def _sub_heading(doc: Document, title: str) -> None:
    """Medium-blue sub-section heading."""
    p = doc.add_paragraph()
    _no_space_para(p)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"  {title}  ")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = _RGB_WHITE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _C_MID)
    pPr.append(shd)


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    _no_space_para(p)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = _RGB_GREY


def _gap(doc: Document, pts: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)


# ── Table builders ────────────────────────────────────────────────────────────

def _kv_table(doc: Document, pairs: list[tuple[str, str]],
              col_w: tuple[float, float] = (7.0, 10.5)) -> None:
    """Two-column key-value table."""
    table = doc.add_table(rows=len(pairs), cols=2)
    _set_table_borders(table)
    for i, (k, v) in enumerate(pairs):
        row = table.rows[i]
        kc, vc = row.cells
        kc.text = k
        vc.text = v
        _set_cell_bg(kc, _C_ALT)
        _no_space_para(kc.paragraphs[0])
        _no_space_para(vc.paragraphs[0])
        for cell in (kc, vc):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
            if cell == kc:
                para.runs[0].font.color.rgb = _RGB_DARK
                para.runs[0].font.bold = True
        if i % 2 == 0:
            _set_cell_bg(vc, _C_WHITE)
        else:
            _set_cell_bg(vc, "F9FAFB")
    _set_col_widths(table, list(col_w))


def _data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    col_w: list[float] | None = None,
    status_cols: list[int] | None = None,   # column indices that contain OK/FAIL/N/A
) -> None:
    """Full-width data table with blue header row."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    _set_table_borders(table)

    # Header row
    hdr = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = h
        _set_cell_bg(cell, _C_DARK)
        _no_space_para(cell.paragraphs[0])
        for run in cell.paragraphs[0].runs:
            run.font.bold  = True
            run.font.size  = Pt(8.5)
            run.font.color.rgb = _RGB_WHITE
            run.font.name  = "Arial"

    status_cols = set(status_cols or [])

    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        bg = _C_ALT if i % 2 == 0 else _C_WHITE
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            txt  = str(val) if val is not None else "—"
            cell.text = txt
            _no_space_para(cell.paragraphs[0])

            # Status-column coloring
            if j in status_cols:
                upper = txt.upper()
                if "✓" in txt or upper.startswith("OK") or upper.startswith("PASS"):
                    _set_cell_bg(cell, _C_OK)
                    cell.paragraphs[0].runs[0].font.color.rgb = _RGB_OK
                elif "✗" in txt or "FAIL" in upper or "ERROR" in upper:
                    _set_cell_bg(cell, _C_FAIL)
                    cell.paragraphs[0].runs[0].font.color.rgb = _RGB_FAIL
                elif "⚠" in txt or "WARN" in upper or "SHORT" in upper or "EXCEED" in upper:
                    _set_cell_bg(cell, _C_WARN)
                    cell.paragraphs[0].runs[0].font.color.rgb = _RGB_WARN
                else:
                    _set_cell_bg(cell, bg)
            else:
                _set_cell_bg(cell, bg)

            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
                    run.font.name = "Arial"

    if col_w:
        _set_col_widths(table, col_w)


def _status(ok) -> str:
    if ok is True:   return "✓ OK"
    if ok is False:  return "✗ FAIL"
    return "— N/A"


# ── Cover page ────────────────────────────────────────────────────────────────

def _cover(doc: Document, project_name: str, vessel_tag: str,
           issued_for: str, today: str, code_full: str,
           gas_fluid: str, liq_fluid: str) -> None:
    # Top blue band
    p = doc.add_paragraph()
    _no_space_para(p)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("  ")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), _C_DARK)
    pPr.append(shd)

    _gap(doc, 18)

    # Title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Horizontal Two-Phase Separator")
    r2.font.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = _RGB_DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Comprehensive Design Report")
    r3.font.size = Pt(14); r3.font.color.rgb = _RGB_MID

    _gap(doc, 14)

    # Info block table
    info = [
        ("Equipment tag",   vessel_tag),
        ("Project",         project_name or "—"),
        ("Service",         f"{gas_fluid} / {liq_fluid}"),
        ("Design code",     code_full),
        ("Issued for",      issued_for),
        ("Date",            today),
        ("Revision",        "A"),
    ]
    _kv_table(doc, info, col_w=(5.5, 12.0))

    _gap(doc, 12)

    # Stamp
    if issued_for != "Construction":
        p_stamp = doc.add_paragraph()
        p_stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space_para(p_stamp)
        r_stamp = p_stamp.add_run(f"  ISSUED FOR {issued_for.upper()} — NOT FOR CONSTRUCTION  ")
        r_stamp.font.bold = True; r_stamp.font.size = Pt(10)
        r_stamp.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        pPr_s = p_stamp._p.get_or_add_pPr()
        shd_s = OxmlElement("w:shd")
        shd_s.set(qn("w:val"), "clear"); shd_s.set(qn("w:fill"), "FEF3C7")
        pPr_s.append(shd_s)

    _gap(doc, 8)
    p_disc = doc.add_paragraph(
        "Generated by SepScope — separator scoping for inquiry and FEED. "
        "Results are screening-level only and must be verified by a qualified "
        "pressure vessel engineer before use in any design or procurement activity."
    )
    p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_disc.runs:
        run.font.size = Pt(8); run.font.color.rgb = _RGB_GREY


# ── Main generator ────────────────────────────────────────────────────────────

def generate_word_report(
    project_name: str,
    vessel_tag: str,
    issued_for: str,
    Di: float,
    L_shell: float,
    h_head: float,
    P_barg: float,
    T_C: float,
    mat_key: str,
    head_type_label: str,
    code_key: str,
    fd_MPa: float,
    shell_res,
    head_res,
    nozzle_results: list,
    levels_mm: dict,
    sep_res,
    gas_props,
    liq_props,
    Q_gas_m3h: float,
    Q_liq_m3h: float,
    gas_fluid: str,
    liq_fluid: str,
    placement_checks: list,
    head_warnings: list,
    shell_warnings: list,
    saddle_a_mm: float,
    saddle_w_mm: float,
    saddle_height_result: dict,
    has_meshpad: bool,
    has_baffles: bool,
    has_inlet_dev: bool,
    has_vortex_brk: bool,
    L_baffle_mm: float,
    baffle_open_pct: float,
    K_sb: float,
    n_inlets: int = 1,
    P_op_barg: float | None = None,
    T_op_C: float | None = None,
    t_holdup_req_min: float = 3.0,
    t_surge_req_min: float = 3.0,
    include_surge_check: bool = True,
    ldv_result: dict | None = None,
    int_loads_result: dict | None = None,
    weight_result: dict | None = None,
    turndown_result: dict | None = None,
    inlet_dev_type: str = "Half-pipe diverter",
    inlet_dev_sizing=None,    # InletDeviceSizing | None
    outlet_vel: dict | None = None,   # keys: dn_go, v_go, rv2_go, rv2_go_ok, dn_lo, v_lo, rv2_lo, rv2_lo_ok
    Z_gas: float = 1.0,
    lining_spec: dict | None = None,
    head_type=None,           # accepted but unused — reserved for future endcap section
    crown_ratio: float = 1.0,
    knuckle_ratio: float = 0.10,
    alpha_deg_cone: float = 30.0,
    ellipse_ratio: float = 2.0,
) -> bytes:
    """Generate a comprehensive Word design report. Returns raw .docx bytes."""
    from engines.nozzle_geometry import NOZZLE_OD, NOZZLE_WALL_T, NOZZLE_WALL_SCH, recommended_schedule
    from engines import MATERIALS
    from engines.head_geometry import HeadType

    today      = _date.today().strftime("%d %b %Y")
    code_full  = "EN 13445-3:2021" if code_key == "EN" else "ASME VIII Div.1"
    pn_label   = "PN" if code_key == "EN" else "Class"
    hydro_f    = 1.25 if code_key == "EN" else 1.43
    hydro_P    = P_barg * hydro_f
    nll_mm     = levels_mm.get("NLL",  Di * 0.5)
    lahh_mm    = levels_mm.get("LAHH", Di * 0.85)
    CA_mm      = getattr(shell_res, "CA_mm", 3.0)
    z_weld     = getattr(shell_res, "z", 1.0)
    mat        = MATERIALS.get(mat_key, {})
    mat_name   = mat.get("name", mat_key)
    rho_g      = gas_props.rho_kgm3
    rho_l      = liq_props.rho_kgm3
    mu_g       = gas_props.mu_Pas * 1e6
    mu_l       = liq_props.mu_Pas * 1e3

    def bore_id(nz_dict):
        dn  = nz_dict["dn"]
        OD  = NOZZLE_OD.get(dn, dn * 1.05)
        rec = recommended_schedule(nz_dict.get("pn", 25), code_key)
        t   = float(NOZZLE_WALL_SCH[rec].get(dn, NOZZLE_WALL_T.get(dn, 8.0)))
        return OD - 2 * t

    doc = Document()
    _setup_document(doc)

    # ── Cover ─────────────────────────────────────────────────────────────────
    _cover(doc, project_name, vessel_tag, issued_for, today,
           code_full, gas_fluid, liq_fluid)
    doc.add_page_break()

    # ── A — Design Conditions ─────────────────────────────────────────────────
    _section_heading(doc, "A", "Design Conditions")
    _kv_table(doc, [
        ("Equipment type",          "Horizontal two-phase separator"),
        ("Orientation",             "Horizontal"),
        ("Design code",             code_full),
        ("Separator standard",      "API 12J"),
        ("Operating pressure",      f"{P_op_barg:.2f} barg" if P_op_barg is not None else "—"),
        ("Operating temperature",   f"{T_op_C:.0f} °C"    if T_op_C  is not None else "—"),
        ("Design pressure",         f"{P_barg:.2f} barg"),
        ("Design temperature",      f"{T_C:.0f} °C"),
        ("Hydrostatic test pressure",
         f"{hydro_P:.2f} barg  ({hydro_f:.2f} × design pressure, {code_full})"),
        ("Fluid service",           "Two-phase gas / liquid — non-fouling"),
        ("Corrosion allowance CA",  f"{CA_mm:.1f} mm"),
        ("Weld joint efficiency z",
         f"{z_weld:.2f}" + (" (full radiography)" if z_weld >= 1.0 else " (partial radiography)")),
    ])

    # ── B — Process Fluids ────────────────────────────────────────────────────
    _section_heading(doc, "B", "Process Fluids")
    _sub_heading(doc, "Gas Phase")
    _kv_table(doc, [
        ("Fluid",                   gas_fluid),
        ("Molecular weight",        f"{gas_props.MW:.2f} g/mol"),
        ("Compressibility Z",       f"{Z_gas:.3f}"),
        ("Density at op. conditions", f"{rho_g:.3f} kg/m³"),
        ("Dynamic viscosity",       f"{mu_g:.2f} μPa·s"),
        ("Volumetric flow rate",    f"{Q_gas_m3h:,.1f} m³/h (actual)"),
        ("Mass flow rate",          f"{Q_gas_m3h * rho_g:,.0f} kg/h"),
    ])
    _sub_heading(doc, "Liquid Phase")
    _kv_table(doc, [
        ("Fluid",                   liq_fluid),
        ("Density at op. conditions", f"{rho_l:.0f} kg/m³"),
        ("Dynamic viscosity",       f"{mu_l:.3f} mPa·s"),
        ("Volumetric flow rate",    f"{Q_liq_m3h:,.2f} m³/h"),
        ("Mass flow rate",          f"{Q_liq_m3h * rho_l:,.0f} kg/h"),
    ])

    # Nozzle stream table
    _sub_heading(doc, "Nozzle Stream Summary")
    inlet_nzs   = [(nz, *r) for nz, *r in nozzle_results if nz.get("service") == "Inlet"]
    gas_out_nzs = [(nz, *r) for nz, *r in nozzle_results if nz.get("service") == "Gas outlet"]
    liq_out_nzs = [(nz, *r) for nz, *r in nozzle_results if nz.get("service") == "Liquid outlet"]
    rho_mix = (rho_g * Q_gas_m3h + rho_l * Q_liq_m3h) / max(Q_gas_m3h + Q_liq_m3h, 1e-9)

    stream_rows = []
    for (nz, *_) in inlet_nzs:
        ID  = bore_id(nz)
        A   = math.pi * (ID * 1e-3) ** 2 / 4.0
        Qg  = Q_gas_m3h / n_inlets
        Ql  = Q_liq_m3h / n_inlets
        Qm  = Qg + Ql
        v_m = (Qm / 3600.0) / max(A, 1e-9)
        rv2 = rho_mix * v_m ** 2
        stream_rows.append([nz["tag"], f"DN{nz['dn']}", nz["service"],
                             "Mixture (two-phase)",
                             f"{Qm:,.1f} m³/h",
                             f"{(Qg*rho_g + Ql*rho_l):,.0f} kg/h",
                             f"{rho_mix:.1f} kg/m³",
                             f"{v_m:.2f} m/s",
                             f"{rv2:,.0f} Pa"])
    for (nz, *_) in gas_out_nzs:
        ID = bore_id(nz)
        A  = math.pi * (ID * 1e-3) ** 2 / 4.0
        v  = (Q_gas_m3h / 3600.0) / max(A, 1e-9)
        stream_rows.append([nz["tag"], f"DN{nz['dn']}", nz["service"],
                             "Gas (separated)",
                             f"{Q_gas_m3h:,.1f} m³/h",
                             f"{Q_gas_m3h*rho_g:,.0f} kg/h",
                             f"{rho_g:.3f} kg/m³",
                             f"{v:.2f} m/s", "—"])
    for (nz, *_) in liq_out_nzs:
        ID = bore_id(nz)
        A  = math.pi * (ID * 1e-3) ** 2 / 4.0
        v  = (Q_liq_m3h / 3600.0) / max(A, 1e-9)
        stream_rows.append([nz["tag"], f"DN{nz['dn']}", nz["service"],
                             "Liquid (separated)",
                             f"{Q_liq_m3h:,.2f} m³/h",
                             f"{Q_liq_m3h*rho_l:,.0f} kg/h",
                             f"{rho_l:.0f} kg/m³",
                             f"{v:.2f} m/s", "—"])
    if stream_rows:
        _data_table(doc,
                    ["Tag", "DN", "Service", "Phase",
                     "Vol. flow", "Mass flow", "Density", "Velocity", "ρv² (Pa)"],
                    stream_rows,
                    col_w=[1.3, 1.2, 2.2, 2.5, 2.0, 2.0, 2.0, 1.8, 1.8])
    _caption(doc, "Inlet ρv² limit: 2 400 Pa (API RP 14E, non-erosive service). "
                  f"Inlet flow split equally between {n_inlets} nozzle(s).")

    # ── C — Separator Sizing ──────────────────────────────────────────────────
    doc.add_page_break()
    _section_heading(doc, "C", "Separator Sizing  (API 12J Screening)")

    delta_rho = max(0.0, rho_l - rho_g)
    pad_rows: list[list] = []
    if has_meshpad:
        U_pad_max = K_sb * math.sqrt(delta_rho / max(rho_g, 0.001))
        A_pad_req = (Q_gas_m3h / 3600.0) / max(U_pad_max, 1e-9)
        A_pad_ok  = A_pad_req <= sep_res.A_gas_m2
        pad_rows.append([
            f"Mesh pad load (full Q_gas)  [K = {K_sb:.2f} m/s]",
            f"{A_pad_req/max(sep_res.A_gas_m2,1e-9)*100:.0f} %  "
            f"(req. {A_pad_req:.3f} m² / avail. {sep_res.A_gas_m2:.3f} m²)",
            "≤ 100 %",
            _status(A_pad_ok),
        ])

    surge_limit = (f"≥ {t_surge_req_min:.1f} min  (required)"
                   if include_surge_check else "Informational only")
    surge_ok_val = sep_res.surge_ok if include_surge_check else None

    sizing_data = [
        ["Slenderness L/D  (T–T / Di)  [API 12J: 3–5]",
         f"{sep_res.LD_ratio:.2f}", "3.0 – 5.0",
         _status(3.0 <= sep_res.LD_ratio <= 5.0)],
        [f"Gas body velocity  [K = {K_sb:.2f} m/s]",
         f"{sep_res.U_act_ms:.3f} m/s", f"≤ {sep_res.U_max_ms:.3f} m/s",
         _status(sep_res.gas_velocity_ok)],
        *pad_rows,
        ["Liquid hold-up time at NLL",
         f"{sep_res.t_holdup_s/60:.1f} min", f"≥ {t_holdup_req_min:.1f} min",
         _status(sep_res.holdup_ok)],
        ["Surge time  NLL → LAHH",
         f"{sep_res.t_surge_s/60:.1f} min", surge_limit,
         _status(surge_ok_val)],
        ["NLL fill fraction  (NLL / Di)  [target 50 %]",
         f"{sep_res.nll_frac*100:.0f} %", "35 – 65 %",
         _status(0.35 <= sep_res.nll_frac <= 0.65)],
        ["Liquid droplet cut size (gas phase)",
         f"{sep_res.d_cut_gas_um:.0f} μm", "Informational", "— N/A"],
        ["Gas bubble cut size (liquid phase)",
         f"{sep_res.d_cut_liq_um:.0f} μm", "Informational", "— N/A"],
    ]
    # Inlet ρv²
    inlet_nzs_raw = [nz for nz, *_ in nozzle_results if nz.get("service") == "Inlet"]
    if inlet_nzs_raw:
        nz0 = inlet_nzs_raw[0]
        ID0 = bore_id(nz0)
        A0  = math.pi * (ID0 * 1e-3) ** 2 / 4.0
        Qm0 = (Q_gas_m3h + Q_liq_m3h) / n_inlets / 3600.0
        v0  = Qm0 / max(A0, 1e-9)
        rv2 = rho_mix * v0 ** 2
        sizing_data.append([
            f"Inlet nozzle ρv²  (DN{nz0['dn']})  [API RP 14E]",
            f"{rv2:,.0f} Pa", "≤ 2 400 Pa", _status(rv2 <= 2400.0),
        ])
    if outlet_vel and outlet_vel.get("rv2_go") is not None:
        sizing_data.append([
            f"Gas outlet ρv²  (DN{outlet_vel['dn_go']})  [API RP 14E C=100 — erosion]",
            f"{outlet_vel['rv2_go']:,.0f} Pa  ({outlet_vel['v_go']:.2f} m/s)",
            f"≤ {outlet_vel.get('rv2_14e', 14884):,.0f} Pa",
            _status(outlet_vel.get("rv2_go_ok")),
        ])
    if outlet_vel and outlet_vel.get("rv2_lo") is not None:
        sizing_data.append([
            f"Liquid outlet ρv²  (DN{outlet_vel['dn_lo']})  [API RP 14E C=100 — erosion]",
            f"{outlet_vel['rv2_lo']:,.0f} Pa  ({outlet_vel['v_lo']:.2f} m/s)",
            f"≤ {outlet_vel.get('rv2_14e', 14884):,.0f} Pa",
            _status(outlet_vel.get("rv2_lo_ok")),
        ])
        sizing_data.append([
            f"Liquid outlet velocity  (DN{outlet_vel['dn_lo']})  [level stability / vortex]",
            f"{outlet_vel['v_lo']:.2f} m/s",
            f"≤ {outlet_vel.get('v_liq_max', 3.0):.1f} m/s",
            _status(outlet_vel.get("v_lo_stab_ok")),
        ])

    _data_table(doc, ["Criterion", "Actual", "Limit / Target", "Status"],
                sizing_data,
                col_w=[8.5, 3.5, 3.5, 2.0],
                status_cols=[3])

    # C.1 — Turndown Analysis
    if turndown_result is not None:
        td = turndown_result
        _sub_heading(doc, f"C.1  Turndown Analysis  ({td['pct']} % of design flow)")
        _caption(doc,
                 f"Service: {td.get('svc_condition', '—')}  ·  "
                 f"Re-entrainment limit = 1.15 × K = {td['U_reentrain_ms']:.3f} m/s  ·  "
                 f"Drainage warning threshold = 0.30 × K limit.")

        def _tds(ok) -> str:
            if ok is True:  return "✓ OK"
            if ok is False: return "✗ FAIL"
            return "—"

        td_rows: list[list] = [[
            f"Gas velocity  (≤ {td['U_max_ms']:.3f} m/s)",
            f"{td['U_act_ms']:.3f} m/s",
            f"{td['U_act_td_ms']:.3f} m/s",
            _tds(td["gas_vel_ok"]),
            _tds(td["gas_vel_td_ok"]),
        ], [
            f"Re-entrainment (≤ {td['U_reentrain_ms']:.3f} m/s, 1.15 × K)",
            f"{td['U_act_ms']:.3f} m/s", "—",
            _tds(td["gas_reentrain_ok"]), "—",
        ]]
        if td.get("has_meshpad") and td.get("pad_load_pct") is not None:
            td_rows.append([
                "Mesh pad load  (≤ 100 %)",
                f"{td['pad_load_pct']:.0f} %", f"{td['pad_td_pct']:.0f} %",
                _tds(td.get("pad_ok")),
                "⚠ drainage" if td.get("pad_drain_warn") else _tds(td.get("pad_td_ok")),
            ])
        td_rows.append([
            f"Hold-up time  (≥ {td['t_holdup_req_min']:.1f} min)",
            f"{td['t_holdup_s']/60:.1f} min", f"{td['t_holdup_td_s']/60:.1f} min",
            _tds(td["holdup_ok"]), _tds(td["holdup_td_ok"]),
        ])
        if td.get("include_surge"):
            td_rows.append([
                f"Surge time NLL→LAHH  (≥ {td['t_surge_req_min']:.1f} min)",
                f"{td['t_surge_s']/60:.1f} min", f"{td['t_surge_td_s']/60:.1f} min",
                _tds(td.get("surge_ok")), _tds(td.get("surge_td_ok")),
            ])
        if td.get("rv2_pa") is not None:
            td_rows.append([
                "Inlet ρv²  (≤ 2 400 Pa)",
                f"{td['rv2_pa']:,.0f} Pa", f"{td['rv2_td_pa']:,.0f} Pa",
                _tds(td.get("rv2_ok")), _tds(td.get("rv2_td_ok")),
            ])
        _data_table(doc,
                    ["Criterion", "Design (100 %)", f"Turndown ({td['pct']} %)",
                     "Design", "TD"],
                    td_rows,
                    col_w=[7.5, 2.5, 2.5, 1.5, 1.5],
                    status_cols=[3, 4])

    # C.2 — LDV
    if ldv_result is not None:
        ldv = ldv_result
        _sub_heading(doc, "C.2  LDV — Liquid Design Volume")
        _caption(doc,
                 "Minimum liquid inventory required to fill downstream equipment that are partially "
                 "empty during operation or startup. Volumes include full vessel geometry "
                 "(cylinder + both endcaps). User specifies required LDV + safety factor; "
                 "two independent checks: Segment A (VB → LZLL) ≥ LDV×SF  and  Segment B (LZLL → LALL) ≥ LDV×SF.")
        _sf_b = ldv.get("sf_b", ldv["sf"])
        ldv_pairs = [
            ("Effective vessel bottom (VB)",
             f"{ldv['eff_vb_mm']:.0f} mm above vessel bottom"),
            ("Safety factor — Seg A", f"{ldv['sf']:.2f}"),
            ("Safety factor — Seg B", f"{_sf_b:.2f}"),
        ]
        if ldv.get("target_m3") is not None:
            _req_a = (ldv.get("ldv_required_a_m3") or ldv.get("ldv_required_m3") or 0.0)
            _req_b = (ldv.get("ldv_required_b_m3") or ldv.get("ldv_required_m3") or 0.0)
            ldv_pairs.append(("Required LDV (before SF)",
                              f"{ldv['target_m3']*1000:.1f} L  ({ldv['target_m3']:.4f} m³)"))
            ldv_pairs.append(("", ""))  # spacer
            ldv_pairs.append(("Segment A (VB → LZLL)",
                              f"{ldv['seg_a_m3']*1000:.1f} L  ({ldv['seg_a_m3']:.4f} m³)"))
            ldv_pairs.append((f"Seg A Required  (Target × SF {ldv['sf']:.2f})",
                              f"{_req_a*1000:.1f} L  ({_req_a:.4f} m³)"))
            ldv_pairs.append(("Segment A ≥ Required?",
                              "✓ PASS" if ldv.get("seg_a_ok") else "✗ FAIL"))
            ldv_pairs.append(("", ""))  # spacer
            ldv_pairs.append(("Segment B (LZLL → LALL)",
                              f"{ldv['seg_b_m3']*1000:.1f} L  ({ldv['seg_b_m3']:.4f} m³)"))
            ldv_pairs.append((f"Seg B Required  (Target × SF {_sf_b:.2f})",
                              f"{_req_b*1000:.1f} L  ({_req_b:.4f} m³)"))
            ldv_pairs.append(("Segment B ≥ Required?",
                              "✓ PASS" if ldv.get("seg_b_ok") else "✗ FAIL"))
            ldv_pairs.append(("", ""))
            ldv_pairs.append(("Both segments adequate?",
                              "✓ PASS" if ldv.get("ok") else "✗ FAIL"))
        else:
            ldv_pairs.append(("Segment A (VB → LZLL)",
                              f"{ldv['seg_a_m3']*1000:.1f} L  ({ldv['seg_a_m3']:.4f} m³)"))
            ldv_pairs.append(("Segment B (LZLL → LALL)",
                              f"{ldv['seg_b_m3']*1000:.1f} L  ({ldv['seg_b_m3']:.4f} m³)"))
            ldv_pairs.append(("Note", "Specify an LDV target to see pass/fail checks."))
        _kv_table(doc, ldv_pairs)

    # C.3 — Internals mechanical loads
    if int_loads_result is not None:
        il = int_loads_result
        _sub_heading(doc, "C.3  Internals — Mechanical Loads (LDV Startup Surge)")
        _caption(doc,
                 f"Governing load case: LDV inventory ({il['V_ldv_m3']*1000:.1f} L) "
                 f"floods into vessel in {il['t_flood_s']:.0f} s, split across "
                 f"{il.get('n_inlets', 1)} inlet(s). "
                 "Pure liquid density assumed (startup slug). "
                 "No standard prescribes this method — verify per project structural code "
                 "(EN 1993-1-8 / AWS D1.1). Safety factors: SF = 3.0 inlet device (impulsive), "
                 "SF = 2.0 baffle (quasi-static).")
        _kv_table(doc, [
            ("LDV surge flow per inlet",
             f"{il['Q_ldv_per_inlet_m3s']*1000:.2f} L/s  "
             f"({il['V_ldv_m3']*1000:.1f} L in {il['t_flood_s']:.0f} s)"),
            ("", ""),
            (f"Inlet device — nozzle DN{il['nozzle_dn']} (ID {il['nz_id_mm']:.0f} mm)",
             f"A = {il['A_nozzle_m2']*1e4:.1f} cm²"),
            ("Inlet surge velocity",   f"{il['v_ldv_ms']:.2f} m/s"),
            ("Impact force (unfactored)", f"{il['F_impact_N']:,.0f} N"),
            (f"Design force  (SF {il['SF_inlet']:.0f})", f"{il['F_inlet_design_N']:,.0f} N"),
            ("Basis", "F = ρ_liq × v² × A_nozzle  (first principles)"),
            ("", ""),
            (f"Baffle plate  (φ = {il['phi']*100:.0f} %, Cd = 0.61)", ""),
            ("Surge ΔP",               f"{il['dP_surge_Pa']:,.0f} Pa"),
            ("Governing force (unfactored)", f"{il['F_baffle_surge_N']:,.0f} N"),
            (f"Design force  (SF {il['SF_baffle']:.0f})", f"{il['F_baffle_design_N']:,.0f} N"),
            ("Gas ΔP operating (ref.)", f"{il['dP_gas_op_Pa']:.1f} Pa"),
            ("", ""),
            ("Min. plate thickness (clamped plate + API 12J ≥ 6 mm)",
             f"{il['t_baffle_design_mm']:.1f} mm  (calc. {il['t_baffle_min_mm']:.1f} mm)"),
            (f"Fillet weld throat  (τ_allow = {il['tau_allow_Pa']/1e6:.0f} MPa = 0.4·f_y)",
             f"{il['a_weld_design_mm']:.1f} mm  (calc. {il['a_weld_req_mm']:.1f} mm, min 3 mm)"),
            ("Weld perimeter", f"{il['L_weld_m']*1000:.0f} mm  (full circumference)"),
            ("Material f_d / f_y", f"{il['fd_MPa']:.0f} MPa / {il['fy_MPa']:.0f} MPa"),
        ])

    # ── D — Mechanical Design ─────────────────────────────────────────────────
    doc.add_page_break()
    _section_heading(doc, "D", "Mechanical Design")
    _kv_table(doc, [
        ("Inner diameter Di",             f"{Di:,.0f} mm"),
        ("Shell length (T–T)",            f"{L_shell:,.0f} mm"),
        ("Overall length (pole–pole)",    f"{L_shell + 2*h_head:,.0f} mm"),
        ("Head type",                     head_type_label),
        ("Head depth h",                  f"{h_head:.0f} mm"),
        ("Material (shell & heads)",      mat_name),
        ("Allowable stress fd",           f"{fd_MPa:.1f} MPa  at {T_C:.0f} °C"),
        ("Corrosion allowance CA",        f"{CA_mm:.1f} mm"),
        ("Weld joint efficiency z",       f"{z_weld:.2f}"),
        ("Shell thickness — calculated",  f"{shell_res.t_calc_mm:.3f} mm"),
        ("Shell thickness — nominal",     f"{shell_res.t_nom_mm:.1f} mm"),
        ("Head thickness — calculated",   f"{head_res.t_calc_mm:.3f} mm"),
        ("Head thickness — nominal",      f"{head_res.t_nom_mm:.1f} mm"),
        ("Support type",                  "Saddle supports — 2 off"),
        ("Saddle position (from tangent)", f"{saddle_a_mm:.0f} mm" if saddle_a_mm > 0 else "TBD"),
        ("Saddle width",                  f"{saddle_w_mm:.0f} mm"),
    ])

    # D.1 — Overall Height & Mounting
    if saddle_height_result is not None:
        sh = saddle_height_result
        _sub_heading(doc, "D.1  Overall Height & Mounting")
        _bn = sh.get("bottom_nozzles", [])
        _bn_txt = (", ".join(f"{b['tag']} DN{b['dn']}" for b in _bn) if _bn else "none")
        _kv_table(doc, [
            ("Outer diameter Dₒ (Di + 2·t)",   f"{sh['Do_mm']:,.0f} mm"),
            ("Saddle stand height",            f"{sh['h_stand_mm']:,.0f} mm"),
            ("Baseplate thickness",            f"{sh['t_base_mm']:.0f} mm"),
            ("Overall height (top → feet)",    f"{sh['overall_height_mm']:,.0f} mm"),
            ("Height basis",                   sh["basis"]),
            ("Governing constraint",           sh["governing"]),
            ("Structural minimum stand",       f"{sh['h_struct_mm']:,.0f} mm"),
            ("Bottom-nozzle clearance req.",
             (f"{sh['h_clear_mm']:,.0f} mm (clears {_bn_txt}, "
              f"{sh['ground_clearance_mm']:.0f} mm ground)") if _bn
             else "— (no bottom nozzles)"),
        ])
        _z = sh.get("zick")
        if _z is not None:
            _bear = (f"{_z['p_act_MPa']:.2f} / {_z['p_allow_MPa']:.1f} MPa  "
                     + ("OK" if _z["bearing_ok"] else "EXCEEDS allowable"))
            _kv_table(doc, [
                ("Saddle reaction (per saddle, hydrotest)",
                 f"{_z['Q_per_saddle_N']/1000:,.0f} kN  ({_z['Q_per_saddle_N']/9.81/1000:.1f} t)"),
                ("Contact (wrap) angle",       f"{_z['wrap_angle_deg']:.0f}°"),
                ("Foundation bearing (actual / allowable)", _bear),
                ("Baseplate B × L × t",
                 f"{_z['B_mm']:,.0f} × {_z['L_bp_mm']:,.0f} × {_z['t_base_mm']:.0f} mm"),
            ])
            _caption(doc, _z["note"])
        for _w in sh.get("warnings", []):
            _caption(doc, f"⚠  {_w}")
        _caption(doc, sh["code_note"])

    # D.2 — Weight estimate
    if weight_result is not None:
        wt = weight_result
        _sub_heading(doc, "D.2  Weight Estimate")
        _caption(doc,
                 "Estimated weights ±15–20 %. Shell and heads use nominal wall thickness. "
                 "Nozzle weight = pipe stub (300 mm projection) + one weld-neck flange per nozzle. "
                 "Saddle weight from plate-area estimate. Misc +5 % covers welds, paint, support clips.")
        _total_w = max(wt["m_dry_kg"], 1.0)
        def _wpct_w(m): return f"{m / _total_w * 100:.1f} %"
        _kv_table(doc, [
            ("Dry weight",
             f"{wt['m_dry_kg']:,.0f} kg  ({wt['m_dry_kg']/1000:.2f} t)"),
            ("Operating weight",
             f"{wt['m_operating_kg']:,.0f} kg  ({wt['m_operating_kg']/1000:.2f} t)  "
             f"[liquid at NLL: {wt['m_liquid_op_kg']:,.0f} kg]"),
            ("Hydrotest weight",
             f"{wt['m_hydrotest_kg']:,.0f} kg  ({wt['m_hydrotest_kg']/1000:.2f} t)  "
             f"[water fill: {wt['m_water_ht_kg']:,.0f} kg]"),
        ])
        _sub_heading(doc, "Dry weight breakdown")
        _data_table(doc,
            ["Component", "Mass (kg)", "% of dry"],
            [
                ["Shell",               f"{wt['m_shell_kg']:,.0f}",    _wpct_w(wt['m_shell_kg'])],
                ["Heads × 2",           f"{wt['m_heads_kg']:,.0f}",    _wpct_w(wt['m_heads_kg'])],
                [f"Nozzles ({len(wt['nozzle_detail'])})",
                                         f"{wt['m_nozzles_kg']:,.0f}", _wpct_w(wt['m_nozzles_kg'])],
                ["Saddles × 2",         f"{wt['m_saddles_kg']:,.0f}",  _wpct_w(wt['m_saddles_kg'])],
                ["Internals",           f"{wt['m_internals_kg']:,.0f}",_wpct_w(wt['m_internals_kg'])],
                [f"Misc (+{wt['misc_factor']*100:.0f} %)",
                                         f"{wt['m_misc_kg']:,.0f}",   f"{wt['misc_factor']*100:.0f} %"],
                ["Total dry",           f"{wt['m_dry_kg']:,.0f}",      "100 %"],
            ],
            col_w=[6.0, 3.5, 3.0],
        )

    _sub_heading(doc, "D.3  Shell Thickness Calculation")
    _kv_table(doc, [
        ("Formula (EN / ASME cylindrical)", shell_res.formula),
        ("Code clause",                     shell_res.clause),
        ("Calculated minimum thickness",    f"{shell_res.t_calc_mm:.3f} mm"),
        ("Corrosion allowance CA",          f"{CA_mm:.1f} mm"),
        ("Nominal thickness selected",      f"{shell_res.t_nom_mm:.1f} mm  (rounded to next 0.5 mm)"),
        ("Weld joint efficiency z",         f"{z_weld:.2f}"),
    ])

    _sub_heading(doc, "D.4  Head Thickness Calculation")
    _kv_table(doc, [
        ("Head type",                  head_type_label),
        ("Formula",                    head_res.formula),
        ("Code clause",                head_res.clause),
        ("Calculated minimum thickness", f"{head_res.t_calc_mm:.3f} mm"),
        ("Corrosion allowance CA",     f"{CA_mm:.1f} mm"),
        ("Nominal thickness selected", f"{head_res.t_nom_mm:.1f} mm"),
    ])
    if head_warnings:
        for w in head_warnings:
            _caption(doc, f"⚠  {w}")
    if shell_res.warnings:
        for w in shell_res.warnings:
            _caption(doc, f"⚠  {w}")

    # D.4 — Internal lining / surface treatment
    if lining_spec:
        ls = lining_spec
        lining_pairs: list[tuple[str, str]] = []
        if ls.get("has_clad"):
            lining_pairs += [
                ("Internal cladding / weld overlay", ls["clad_material"]),
                ("Cladding thickness", f"{ls['clad_t_mm']:.1f} mm  (min., after forming)"),
                ("Note", "Cladding does not contribute to pressure-bearing wall thickness"),
            ]
            if ls.get("clad_note"):
                lining_pairs.append(("Cladding specification", ls["clad_note"]))
        if ls.get("has_enp"):
            lining_pairs += [
                ("Internal surface plating", ls["enp_type"]),
                ("Plating thickness", f"{ls['enp_t_um']:.0f} µm  (min.)"),
                ("Note", "Plating does not contribute to pressure-bearing wall thickness"),
            ]
            if ls.get("enp_note"):
                lining_pairs.append(("Plating specification", ls["enp_note"]))
        if ls.get("free_text"):
            lining_pairs.append(("Additional material / treatment notes", ls["free_text"]))
        if lining_pairs:
            _sub_heading(doc, "D.5  Internal Lining / Surface Treatment")
            _kv_table(doc, lining_pairs)

    # ── E — Liquid Levels ─────────────────────────────────────────────────────
    _section_heading(doc, "E", "Liquid Levels & Volumes")
    _level_order = ["LZLL", "LALL", "LAL", "NLL", "LAH", "LAHH", "LZHH"]
    _level_desc  = {
        "LZLL": "Low-low liquid level — low-low shutdown trip",
        "LALL": "Low alarm liquid level",
        "LAL":  "Low liquid level — lower operating bound",
        "NLL":  "Normal liquid level — design basis",
        "LAH":  "High liquid level — upper operating bound",
        "LAHH": "High-high liquid level — high-high shutdown / surge basis",
        "LZHH": "High-high-high liquid level — overfill / trip",
    }
    from engines.vessel_volume import vessel_volumes as _vessel_volumes
    from engines.separator_process import _cyl_vol_mm3
    V_total = sep_res.V_total_vessel_m3 if sep_res.V_total_vessel_m3 > 0 else 1.0
    _lvl_tags_w = [t for t in _level_order if t in levels_mm]
    if head_type is not None and _lvl_tags_w:
        _vr_w = _vessel_volumes(
            head_type, Di, L_shell,
            {t: levels_mm[t] for t in _lvl_tags_w},
            crown_ratio=crown_ratio, knuckle_ratio=knuckle_ratio,
            alpha_deg_cone=alpha_deg_cone, ellipse_ratio=ellipse_ratio,
            include_heads=True,
        )
        _lvols_w = {r["tag"]: r["vol_m3"] for r in _vr_w["levels"]}
    else:
        _lvols_w = {}

    level_rows = []
    for tag in _level_order:
        if tag not in levels_mm:
            continue
        h = max(0.0, min(Di, levels_mm[tag]))
        vol = _lvols_w.get(tag, _cyl_vol_mm3(Di, L_shell, h) * 1e-9)
        level_rows.append([
            tag, _level_desc.get(tag, ""),
            f"{h:.0f}", f"{h/Di*100:.1f}",
            f"{vol:.3f}", f"{vol*1000:.0f}",
        ])
    _data_table(doc,
                ["Tag", "Description", "Height (mm)", "% Di",
                 "Volume (m³)", "Volume (L)"],
                level_rows,
                col_w=[1.5, 6.5, 2.0, 1.5, 3.0, 2.0])

    # ── F — Internals ─────────────────────────────────────────────────────────
    def _word_dev_desc(dev_type: str, has_dev: bool, ids) -> str:
        if not has_dev:
            return "None"
        if dev_type == "Vane distributor (vendor-sized)":
            return "Vane distributor — size per vendor data sheet"
        if ids is None:
            return dev_type
        if ids.device_type == "Half-pipe diverter":
            return (f"Half-pipe diverter  OD {ids.D_device_mm:.0f} mm × "
                    f"L {ids.L_device_mm:.0f} mm  |  "
                    f"Face {ids.A_opening_m2*1e4:.0f} cm² ({ids.area_ratio:.1f}× nozzle)  |  "
                    f"ρv² {ids.rv2_face_Pa:,.0f}/{ids.rv2_limit_Pa:,.0f} Pa "
                    f"({'OK' if ids.adequate else 'FAIL'})")
        if ids.device_type == "Slotted/perforated cylinder":
            return (f"Perforated cylinder  OD {ids.D_device_mm:.0f} mm × "
                    f"L {ids.L_device_mm:.0f} mm  |  "
                    f"Slots {ids.A_slot_mm2:,.0f} mm² ({ids.area_ratio:.1f}× nozzle)  |  "
                    f"ρv² {ids.rv2_face_Pa:,.0f}/{ids.rv2_limit_Pa:,.0f} Pa "
                    f"({'OK' if ids.adequate else 'FAIL'})")
        return dev_type

    _section_heading(doc, "F", "Internals")
    _data_table(doc,
                ["Component", "Description", "Qty", "Function / Notes"],
                [
                    ["Inlet device",
                     _word_dev_desc(inlet_dev_type, has_inlet_dev, inlet_dev_sizing),
                     str(n_inlets) if has_inlet_dev else "—",
                     "Deflects two-phase flow downward; reduces jetting and surface turbulence"],
                    ["Inlet baffles",
                     (f"Perforated plate {baffle_open_pct:.0f} % open, "
                      f"setback {L_baffle_mm:.0f} mm from each tangent") if has_baffles else "None",
                     "2" if has_baffles else "—",
                     f"Uniform flow distribution; effective separation length "
                     f"= {max(0.0, L_shell-2*L_baffle_mm):.0f} mm"],
                    ["Gas demister",
                     f"Knitted wire mesh pad, K_pad = {K_sb:.2f} m/s" if has_meshpad else "None",
                     "1" if has_meshpad else "—",
                     "Removes entrained liquid droplets from gas; located upstream of gas outlet"],
                    ["Vortex breaker",
                     "Cross-plate vortex breaker" if has_vortex_brk else "None",
                     "1" if has_vortex_brk else "—",
                     "Fitted at liquid outlet; prevents gas core formation at low liquid levels"],
                ],
                col_w=[2.8, 5.0, 1.0, 8.7])

    # F.1 — Inlet Device Sizing
    if inlet_dev_sizing is not None:
        ids = inlet_dev_sizing
        _sub_heading(doc, "F.1  Inlet Device Sizing  (API 12J §5.3)")
        if ids.device_type == "Half-pipe diverter":
            _kv_table(doc, [
                ("Device type",          "Half-pipe diverter  (API 12J §5.3.1)"),
                ("Half-pipe OD",         f"{ids.D_device_mm:.0f} mm  (≥ 1.5 × nozzle OD {ids.nozzle_OD_mm:.0f} mm)"),
                ("Half-pipe length",     f"{ids.L_device_mm:.0f} mm"),
                ("Face area",            f"{ids.A_opening_m2*1e4:.1f} cm²"),
                ("Area ratio  (A_face / A_nozzle)", f"{ids.area_ratio:.2f}  (≥ 2.0 required)"),
                ("Impact velocity",      f"{ids.v_face_ms:.2f} m/s"),
                ("ρv² at device face",   f"{ids.rv2_face_Pa:,.0f} Pa  (limit {ids.rv2_limit_Pa:,.0f} Pa)"),
                ("Overall adequacy",     "✓ OK" if ids.adequate else "✗ FAIL"),
            ])
        elif ids.device_type == "Slotted/perforated cylinder":
            _kv_table(doc, [
                ("Device type",         "Slotted/perforated cylinder  (API 12J §5.3.2)"),
                ("Cylinder OD",         f"{ids.D_device_mm:.0f} mm"),
                ("Cylinder length",     f"{ids.L_device_mm:.0f} mm"),
                ("Total slot area",     f"{ids.A_slot_mm2:,.0f} mm²  ({ids.area_ratio:.2f} × nozzle area)"),
                ("Indicative holes",    f"{ids.n_holes_dn25} × DN25"),
                ("Velocity through slots", f"{ids.v_face_ms:.2f} m/s"),
                ("ρv² at slots",        f"{ids.rv2_face_Pa:,.0f} Pa  (limit {ids.rv2_limit_Pa:,.0f} Pa)"),
                ("Overall adequacy",    "✓ OK" if ids.adequate else "✗ FAIL"),
            ])
        elif ids.device_type == "Vane distributor (vendor-sized)":
            _kv_table(doc, [("Device type", "Vane distributor"), ("Sizing", "Per vendor data sheet")])

    # ── G — Nozzle Schedule ───────────────────────────────────────────────────
    doc.add_page_break()
    _section_heading(doc, "G", "Nozzle Schedule")
    lzhh_mm = levels_mm.get("LZHH", 0.0)
    nz_rows = []
    for nz, nres, rres, fok, pat in nozzle_results:
        dn   = nz["dn"]
        OD   = NOZZLE_OD.get(dn, dn * 1.05)
        rec  = recommended_schedule(nz.get("pn", 25), code_key)
        t    = float(NOZZLE_WALL_SCH[rec].get(dn, NOZZLE_WALL_T.get(dn, 8.0)))
        geom_ok  = nres.geom_ok  if nres else True
        code_ok  = nres.code_ok  if nres else True
        reinf_ok = rres.adequate if rres else True
        all_ok   = geom_ok and (code_ok is not False) and fok and (reinf_ok is not False)
        ok_flag  = None if (all_ok and code_ok is None) else all_ok
        notes    = []
        if nres:
            notes.append(f"Zone: {nres.zone.replace('_', ' ')}")
            if nz.get("loc") in ("Left head", "Right head"):
                nz_IR_w  = (OD - 2*t) / 2.0
                nz_bot_w = (Di - nres.d_from_top_mm) - nz_IR_w
                top_clr_w  = nres.edge_to_shell_mm
                lzhh_clr_w = nz_bot_w - lzhh_mm
                notes.append(f"OD top→crown: {top_clr_w:.0f} mm")
                if nz.get("service") == "Inlet":
                    _fw = ("✓" if lzhh_clr_w >= 150
                           else ("✗ sub." if lzhh_clr_w < 0 else "⚠ <150mm"))
                    notes.append(f"LZHH→inlet bot: {lzhh_clr_w:.0f} mm {_fw}")
        # Schedule upgrade recommendation
        if rres is not None and not rres.adequate:
            _t_req_w = shell_res.t_calc_mm if nres is None else head_res.t_calc_mm
            _t_nom_w = shell_res.t_nom_mm  if nres is None else head_res.t_nom_mm
            _upg_w = suggest_schedule_upgrade(
                Di=Di, P_barg=P_barg, fd_MPa=fd_MPa,
                nozzle_OD_mm=OD, current_schedule=rec, nozzle_dn=dn,
                t_req_mm=_t_req_w, t_nom_mm=_t_nom_w,
                CA_mm=CA_mm, code=code_key, z=z_weld,
                space_to_wall_mm=nres.edge_to_shell_mm if nres else None,
                space_to_knuckle_mm=nres.edge_to_knuckle_mm if nres else None,
            )
            if _upg_w:
                notes.append(f"Reinf. fail → {_upg_w}")
        nz_rows.append([
            nz["tag"], "1", nz["service"], nz["loc"],
            f"DN{dn}", f"{pn_label} {nz.get('pn','')}",
            f"{OD:.1f}", f"{t:.1f}", rec, "RF",
            _status(ok_flag), "  |  ".join(notes),
        ])
    _data_table(doc,
                ["Tag", "Qty", "Service", "Location", "DN", pn_label,
                 "OD mm", "Wall mm", "Sch.", "Facing", "Status", "Notes"],
                nz_rows,
                col_w=[1.2, 0.7, 2.2, 2.5, 1.2, 1.2, 1.3, 1.5, 1.5, 1.2, 1.5, 2.5],
                status_cols=[10])

    # G.1 — Endcap nozzle detail
    head_nz = [(nz, nres, rres) for nz, nres, rres, *_ in nozzle_results
               if nres is not None and nz.get("loc") in ("Left head", "Right head")]
    if head_nz:
        _sub_heading(doc, "G.1  Endcap Nozzle Placement & Inlet Positioning")
        R = Di / 2.0
        min_weld_clr = max(3.0 * head_res.t_nom_mm, 25.0)

        detail_rows = []
        for nz, nres, rres in head_nz:
            r        = nres.r_from_axis_mm
            e2k      = f"{nres.edge_to_knuckle_mm:.0f} mm" if nres.edge_to_knuckle_mm is not None else "—"
            geom_s   = _status(nres.geom_ok)
            code_s   = _status(nres.code_ok)
            reinf_s  = _status(rres.adequate if rres else None)
            weld_ok  = nres.edge_to_shell_mm >= min_weld_clr
            top_clr  = nres.edge_to_shell_mm
            nz_IR_d  = (nres.nozzle_OD_mm - 2 * nres.nozzle_t_mm) / 2.0
            nz_bot_d = (Di - nres.d_from_top_mm) - nz_IR_d
            lzhh_clr = nz_bot_d - lzhh_mm
            _lf = ("✓" if lzhh_clr >= 150
                   else ("✗ sub" if lzhh_clr < 0 else "⚠<150"))
            inlet_s  = (f"{lzhh_clr:.0f} mm {_lf}"
                        if nz.get("service") == "Inlet" else "—")
            detail_rows.append([
                nz["tag"], nz["loc"],
                f"DN{nz['dn']}  ({nres.nozzle_OD_mm:.1f})",
                f"{nres.d_from_top_mm:.0f} mm",
                f"{r:.0f}  ({r/R:.2f}×R)",
                nres.zone.replace("_", " ").capitalize(),
                f"{top_clr:.0f} mm  " + ("✓" if weld_ok else "✗"),
                e2k,
                f"{nres.z_on_head_mm:.0f} / {nres.head_depth_mm:.0f} mm",
                inlet_s,
                geom_s, code_s, reinf_s,
            ])
        _data_table(doc,
                    ["Tag", "Side", "DN (OD mm)", "From top", "r (r/R)",
                     "Zone", "OD top→crown", "Edge→bnd",
                     "Depth/Head", "LZHH→inlet bot",
                     "Geom", "Code", "Reinf"],
                    detail_rows,
                    col_w=[1.1, 1.8, 2.1, 1.6, 2.0, 2.0, 2.2, 1.8, 2.0, 2.0, 1.1, 1.1, 1.1],
                    status_cols=[10, 11, 12])

        # Inlet positioning detail for inlet nozzles
        inlet_head_nz = [(nz, nres) for nz, nres, _ in head_nz
                         if nz.get("service") == "Inlet"]
        if inlet_head_nz:
            _sub_heading(doc, "G.1.1  Inlet Nozzle Positioning Detail")
            for nz, nres in inlet_head_nz:
                nz_IR_p  = (nres.nozzle_OD_mm - 2 * nres.nozzle_t_mm) / 2.0
                nz_bot_p = (Di - nres.d_from_top_mm) - nz_IR_p
                top_clr_p  = nres.edge_to_shell_mm
                lzhh_clr_p = nz_bot_p - lzhh_mm
                _kv_table(doc, [
                    (f"{nz['tag']} — {nz['service']} ({nz['loc']}, DN{nz['dn']})", ""),
                    ("Nozzle OD top → vessel crown ID",
                     f"{top_clr_p:.0f} mm"
                     + ("  ✓" if top_clr_p >= min_weld_clr
                        else f"  ✗ below min {min_weld_clr:.0f} mm")),
                    ("LZHH → inlet device bottom",
                     f"{lzhh_clr_p:.0f} mm"
                     + ("  ✓ ≥ 150 mm" if lzhh_clr_p >= 150
                        else ("  ✗ SUBMERGED at LZHH" if lzhh_clr_p < 0
                              else f"  ⚠ only {lzhh_clr_p:.0f} mm — min 150 mm"))),
                    ("Inlet bore bottom from vessel bottom", f"{nz_bot_p:.0f} mm"),
                    ("Nozzle centreline from vessel top", f"{nres.d_from_top_mm:.0f} mm"),
                ], col_w=(7.0, 10.5))

        # Nozzle warnings per head nozzle
        for nz, nres, rres in head_nz:
            all_msgs = (nres.errors + nres.warnings) + (rres.warnings if rres else [])
            if all_msgs:
                _caption(doc, f"{nz['tag']}: " + "  |  ".join(all_msgs[:3]))

    # ── H — Engineering Findings & Notes ─────────────────────────────────────
    _section_heading(doc, "H", "Engineering Findings & Notes")

    _sub_heading(doc, "H.1  Engineering Findings")
    all_findings = []

    # Sizing-table failures (Section C) — consolidate so H is not "no issues"
    # while the sizing table shows a failed criterion.
    for _r in sizing_data:
        if len(_r) >= 4 and "FAIL" in str(_r[3]):
            _crit_clean = str(_r[0]).split("  [")[0]
            all_findings.append(
                ("error", "[Sizing] ", f"{_crit_clean}: {_r[1]}  (limit {_r[2]})"))

    # Nozzle reinforcement / geometry failures (Section G)
    for _nz, _nres, _rres, *_ in nozzle_results:
        if _rres is not None and _rres.adequate is False:
            if _rres.A_deficit_mm2 > 0:
                all_findings.append(
                    ("error", "[Nozzle] ",
                     f"{_nz['tag']} reinforcement inadequate "
                     f"(area deficit {_rres.A_deficit_mm2:,.0f} mm²) — pad or schedule upgrade required."))
            elif _rres.warnings:
                all_findings.append(("error", "[Nozzle] ", f"{_nz['tag']}: {_rres.warnings[0]}"))
            else:
                all_findings.append(
                    ("error", "[Nozzle] ",
                     f"{_nz['tag']} reinforcement not valid as placed — specialist analysis required."))
        if _nres is not None and _nres.geom_ok is False:
            all_findings.append(
                ("error", "[Nozzle] ",
                 f"{_nz['tag']} geometry not buildable as placed — relocate or resize."))

    for w in head_warnings + shell_res.warnings:
        all_findings.append(("warning", "", w))
    for chk in placement_checks:
        tags = f"[{', '.join(chk.tags)}] " if chk.tags else ""
        all_findings.append((chk.level, tags, chk.headline))

    if all_findings:
        find_rows = []
        for level, tags, msg in all_findings:
            icon = {"error": "✗ Error", "warning": "⚠ Warning", "info": "ℹ Info"}.get(level, "")
            find_rows.append([icon, f"{tags}{msg}"])
        _data_table(doc, ["Severity", "Finding"], find_rows,
                    col_w=[2.5, 15.0], status_cols=[0])
    else:
        _caption(doc, "✓  No engineering issues detected.")

    _sub_heading(doc, "H.2  Placement Check Detail")
    for chk in placement_checks:
        tags_str = f"[{', '.join(chk.tags)}] " if chk.tags else ""
        p = doc.add_paragraph()
        _no_space_para(p)
        p.paragraph_format.space_before = Pt(4)
        icon = {"error": "✗", "warning": "⚠", "info": "ℹ"}.get(chk.level, "")
        run = p.add_run(f"{icon}  {tags_str}{chk.headline}")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = {"error": _RGB_FAIL, "warning": _RGB_WARN}.get(chk.level, _RGB_GREY)
        if chk.code_ref:
            p2 = doc.add_paragraph(f"    Code ref: {chk.code_ref}")
            _no_space_para(p2)
            for r2 in p2.runs:
                r2.font.size = Pt(8.5); r2.font.color.rgb = _RGB_GREY
        if chk.detail:
            p3 = doc.add_paragraph(f"    {chk.detail}")
            _no_space_para(p3)
            for r3 in p3.runs:
                r3.font.size = Pt(8.5)
        if chk.impact:
            p4 = doc.add_paragraph(f"    Action: {chk.impact}")
            _no_space_para(p4)
            for r4 in p4.runs:
                r4.font.size = Pt(8.5); r4.font.italic = True

    _sub_heading(doc, "H.3  General Notes")
    notes = [
        "All dimensions in millimetres unless otherwise stated.",
        f"Design code: {code_full}. Separator sizing standard: API 12J.",
        f"Hydrostatic test pressure: {hydro_P:.2f} barg ({hydro_f:.2f} × design pressure, water fill).",
        "Vessel weights (empty / operating / hydro test) to be confirmed by the fabricator.",
        "Nozzle schedule is indicative for inquiry; final schedule to be confirmed by the process engineer.",
        "All calculations are screening-level. Final design must be verified by a qualified "
        "pressure vessel engineer in accordance with the applicable design code.",
        "This document is generated by SepScope and is not a certified engineering calculation.",
    ]
    for i, note in enumerate(notes, 1):
        p_n = doc.add_paragraph(f"  {i}.  {note}")
        _no_space_para(p_n)
        p_n.paragraph_format.space_before = Pt(2)
        for run in p_n.runs:
            run.font.size = Pt(8.5)

    # ── Save ──────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
